import PyPDF2
import docx
from googletrans import Translator
pdfFileObj = open('c:/Users/user_name/Location of PDF/PDF2Doc.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
pdfReader.numPages
doc = docx.Document()
doc.save('c:/Users/user_name/Location for new document/New.docx')
for pageNum in range(0, pdfReader.numPages):
    pageObj = pdfReader.getPage(pageNum)
    pageObj.extractText()
    translator = Translator()
    translations = translator.translate([pageObj.extractText()], dest='en')
    for translation in translations:
         print(translation.text)
    doc = docx.Document('c:/Users/user_name/Location for new document/New.docx')
    doc.add_paragraph(translation.text)
    doc.save('c:/Users/user_name/Location for new document/New.docx')
doc.save('c:/Users/user_name/Location for new document/New.docx')
input("Press Enter to exit")
quit()
